from __future__ import annotations

import json
import os
from typing import Any, Dict, List, Optional

import pandas as pd
from dotenv import load_dotenv

from .dossier import Dossier
from .sql_templates import TEMPLATE_REGISTRY

load_dotenv()

# --- 依赖库按需导入 (防止用户没装包报错) ---
try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from quantchdb import ClickHouseDatabase
except ImportError:
    ClickHouseDatabase = None

class DualModeLoader:
    """
    【双模加载器】(Dual Mode Loader)  
    这是系统的“数据入口”。它支持三种模式：
    1) 本地文件夹模式 (Local Mode)
    2) 数据库模式 (ClickHouse Mode)
    3) API 模式 (API Mode) - 预留扩展
    """

    DEFAULT_TABLE_NAME_MAP: Dict[str, str] = {
        "sampled_etf_basic": "etf_basic",
        "govcn_2025": "govcn",
        "etf_2025_data": "etf_daily",
    }

    def __init__(self):
        pass

    # ================= 模式 A: 本地文件 (保持不变) =================
    def load_from_folder(self, 
                         mission: str, 
                         folder_path: str, 
                         file_map: Optional[Dict[str, str]] = None,
                         table_name_map: Optional[Dict[str, str]] = None,          
                         table_name_map_path: Optional[str] = None,               
                         auto_load_table_map_json: bool = True, 
                         ) -> Dossier:
        """
        扫描指定文件夹，自动识别并加载所有支持的文件。
        
        支持格式：
        - .csv  -> Table
        - .xlsx -> 多 Sheet -> 多 Table
        - .txt/.md/.docx/.pdf -> Text
        """
        dossier = Dossier.create_empty(mission=mission)
        dossier.meta["source_path"] = folder_path

        file_map = file_map or {}
        table_name_map = table_name_map or {}

        if not os.path.exists(folder_path):
            print(f"❌ [Loader] 错误：路径不存在 -> {folder_path}")
            return dossier

        # 外部映射：默认尝试读取 folder/table_map.json
        external_map: Dict[str, str] = {}
        if auto_load_table_map_json and not table_name_map_path:
            candidate = os.path.join(folder_path, "table_map.json")
            if os.path.exists(candidate):
                table_name_map_path = candidate

        if table_name_map_path and os.path.exists(table_name_map_path):
            try:
                with open(table_name_map_path, "r", encoding="utf-8") as f:
                    obj = json.load(f)
                if isinstance(obj, dict):
                    external_map = {str(k): str(v) for k, v in obj.items()}
                    print(f"   ✅ [Loader] 已加载外部表名映射: {table_name_map_path}")
            except Exception as e:
                print(f"   ⚠️ [Loader] 外部映射读取失败 {table_name_map_path}: {e}")

        # 合并映射优先级：默认 < 外部 < 运行时
        merged_map: Dict[str, str] = {}
        merged_map.update(self.DEFAULT_TABLE_NAME_MAP)
        merged_map.update(external_map)
        merged_map.update(table_name_map)

        # 把 merged_map 也注册到 dossier.alias 系统
        dossier.register_table_aliases(merged_map)

        print(f"📂 [Loader] 正在扫描文件夹: {folder_path} ...")
        if file_map:
            print(f"   (启用文件名映射 file_map: {file_map})")
        if merged_map:
            print(f"   (启用表名映射 table_name_map: {merged_map})")

        for filename in os.listdir(folder_path):
            if filename.startswith("."):
                continue

            file_path = os.path.join(folder_path, filename)
            fname_lower = filename.lower()
            base_name = os.path.splitext(filename)[0]

            # 先 file_map（精确文件名），否则用 base_name
            target_name = file_map.get(filename, base_name)

            # 再 base_name -> canonical（两步兜底）
            target_name = merged_map.get(target_name, target_name)
            target_name = merged_map.get(base_name, target_name)

            # --- 结构化数据 ---
            if fname_lower.endswith(".csv"):
                self._load_csv(dossier, file_path, target_name, aliases=[base_name])  # ✅【修改点 L10】登记 alias
            elif fname_lower.endswith(".xlsx"):
                self._load_excel(dossier, file_path, target_name)
            # --- 非结构化文本 ---
            elif fname_lower.endswith((".txt", ".md")):
                self._load_txt(dossier, file_path, filename)
            elif fname_lower.endswith(".docx"):
                self._load_docx(dossier, file_path, filename)
            elif fname_lower.endswith(".pdf"):
                self._load_pdf(dossier, file_path, filename)

        print(f"✅ [Loader] 加载完成。")
        return dossier
    
    # ================= 模式 B: 数据库集成 (ClickHouse) =================
    def load_from_clickhouse(self, 
                             mission: str, 
                             # --- 核心参数 ---
                             sql: Optional[str] = None,
                             template_name: Optional[str] = None,
                             
                             # --- 案卷参数 ---
                             table_name_in_dossier: str = "db_result",
                             
                             # --- 连接参数 (通常从 .env 读，不用传) ---
                             host: Optional[str] = None, 
                             port: Optional[int] = None, 
                             user: Optional[str] = None, 
                             password: Optional[str] = None,
                             database: Optional[str] = None,
                             
                             # --- 模版动态参数 (关键) ---
                             **kwargs) -> Dossier:
        """
        [数据库模式] 执行 SQL 从 ClickHouse 获取数据。
        [数据库通用入口]
        用法 1：直接 SQL
        用法 2：template_name + kwargs
        用法 3：kwargs 里传 table_name -> 自动 universal 模版
        """
        dossier = Dossier.create_empty(mission=mission)
        dossier.meta["source_type"] = "clickhouse_tcp"

        # --- 1. 逻辑分流：决定到底执行哪句 SQL ---
        final_sql = ""
        # 情况 A: 用户直接给了 SQL -> 听用户的
        if sql:
            final_sql = sql 
        # 情况 B: 用户给了模版名 -> 查字典生成
        elif template_name:
            if template_name not in TEMPLATE_REGISTRY:
                print(f"❌ [Loader] 找不到模版: {template_name}")
                return dossier
            try:
                final_sql = TEMPLATE_REGISTRY[template_name](**kwargs)
            except Exception as e:
                print(f"❌ [Loader] 模版生成出错: {e}")
                return dossier
        # 情况 C: 用户啥都没给，但 kwargs 里有 'table_name' -> 自动启用万能模版
        elif "table_name" in kwargs:
            print(f"ℹ️ [Loader] 检测到 table_name，自动启用万能模版...")
            final_sql = TEMPLATE_REGISTRY["universal"](**kwargs)  
        else:
            print("❌ [Loader] 必须提供 sql, template_name 或 table_name 其中之一")
            return dossier

        dossier.meta["sql"] = final_sql
        print(f"🔧 [Loader] 准备执行 SQL: {final_sql[:100]}...")

        # --- 2. 建立连接与执行 (标准流程) ---
        if ClickHouseDatabase is None:
            print("❌ [Loader] 缺少 quantchdb 库，请确保已安装。")
            return dossier

        _host = host or os.getenv("CLICKHOUSE_HOST", "localhost")
        _port = port or int(os.getenv("CLICKHOUSE_PORT", "8123"))
        _user = user or os.getenv("CLICKHOUSE_USER", "default")
        _password = password or os.getenv("CLICKHOUSE_PASSWORD", "")
        _database = database or os.getenv("CLICKHOUSE_DB", "default")

        try:
            print(f"🔌 [Loader] 连接数据库 ({_host})...")
            db = ClickHouseDatabase(
                config={
                    "host": _host,
                    "port": _port,
                    "user": _user,
                    "password": _password,
                    "database": _database,
                },
                terminal_log=False,
                file_log=False,
            )
            raw_data = db.fetch(final_sql)
            df = pd.DataFrame(raw_data)

            # --- 4. 智能表头优化 (Smart Columns) ---
            req_cols= kwargs.get('columns')
            if req_cols and isinstance(req_cols, list) and len(df.columns) == len(req_cols):
                df.columns = req_cols
                print(f"   -> 已自动匹配列名: {req_cols}")
                    
            final_table_name = kwargs.get("table_name", table_name_in_dossier)
            dossier.add_table(
                name=final_table_name,
                df=df,
                description=f"Source: DB ({len(df)} rows)",
                source="clickhouse",
            )
            print(f"✅ [Loader] 成功获取 {len(df)} 行数据 -> 表名: {final_table_name}")


        except Exception as e:
            print(f"⚠️ [Loader] 数据库查询失败: {e}")

        return dossier


    # ================== 模式 C: API 生态扩展 ==================
    def load_from_api(self, mission: str, api_data: Dict[str, Any]) -> Dossier:
        """[生态模式] 预留接口，供业务人员扩展。"""
        dossier = Dossier.create_empty(mission=mission)
        dossier.meta["source_type"] = "api_integration"
        print("⚠️ [Loader] API 模式尚未实现具体的解析逻辑 (TODO: 业务人员可在此处扩展)")
        return dossier
    

    # ================= 辅助功能: 查看表结构 =================
    def inspect_table(self, table_name: str) -> List[str]:
        """[探路功能] 返回表的列名列表"""
        check_sql = f"SELECT * FROM {table_name} LIMIT 1"
        temp = self.load_from_clickhouse(mission="inspect", sql=check_sql)
        if temp.structured_data:
            df = list(temp.structured_data.values())[0]
            cols = list(df.columns)
            print(f"👀 [Inspector] 表 '{table_name}' 包含: {cols}")
            return cols
        return []
    
    # ================= 内部处理逻辑 (Private Methods) =================
    def _load_csv(self, dossier: Dossier, path: str, table_name: str, aliases: Optional[List[str]] = None) -> None:
        encodings_to_try = ["utf-8-sig", "utf-8", "gb18030", "latin1"]
        last_err = None

        for enc in encodings_to_try:
            try:
                df = pd.read_csv(path, encoding=enc)
                df.columns = [str(c).strip() for c in df.columns]
                dossier.add_table(
                    name=table_name,
                    df=df,
                    description=f"CSV Source (encoding={enc})",
                    source=path,
                    aliases=aliases,
                )
                print(f"  -> 已加载表: {table_name} ({len(df)} rows, encoding={enc})")
                return
            except Exception as e:
                last_err = e

        print(f"  ⚠️ CSV读取失败 {path}: {last_err}")

    def _load_excel(self, dossier: Dossier, path: str, base_name: str):
        try:
            dfs = pd.read_excel(path, sheet_name=None)
            for sheet_name, df in dfs.items():
                full_key = base_name if len(dfs) == 1 else f"{base_name}_{sheet_name}"
                dossier.add_table(name=full_key, df=df, description=f"Excel Source")
        except Exception as e:
            print(f"  ⚠️ Excel读取失败 {path}: {e}")

    def _load_txt(self, dossier: Dossier, path: str, filename: str):
        try:
            with open(path, "r", encoding="utf-8") as f:
                dossier.add_text(content=f.read(), source=filename)
        except Exception as e:
            print(f"  ⚠️ 文本读取失败 {filename}: {e}")

    def _load_docx(self, dossier: Dossier, path: str, filename: str):
        if Document is None:
            print(f"  ⚠️ 缺少 docx 库，跳过: {filename}")
            return
        
        try:
            doc = Document(path)
            full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            if full_text:
                dossier.add_text(content=full_text, source=filename)
                print(f"  [Loader] 已提取 Word: {filename}")
        except Exception as e:
            print(f"  ⚠️ Word读取失败 {filename}: {e}")

    def _load_pdf(self, dossier: Dossier, path: str, filename: str):
        if PdfReader is None:
            print(f"  ⚠️ 缺少 pypdf 库，跳过: {filename}")
            return

        try:
            reader = PdfReader(path)
            full_text = "\n".join([p.extract_text() for p in reader.pages if p.extract_text()])
            if full_text:
                dossier.add_text(content=full_text, source=filename)
                print(f"  [Loader] 已提取 PDF: {filename}")
        except Exception as e:
            print(f"  ⚠️ PDF读取失败 {filename}: {e}")
